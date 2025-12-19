import json
import re
from typing import Any, Dict, Iterable, List, Optional, Union

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

DOCX_PATH = "/Users/afschowdhury/Code Local/itac-report-validator/docs/report1/LS2502 - Final Draft R2.docx"

# ---------- Low-level helpers ----------
from icecream import ic
ic.configureOutput(includeContext=True, prefix='DEBUG: ')

# Import link extractor
from link_extractor import extract_links_from_all_ars, get_link_statistics

def iter_block_items(doc: Document) -> Iterable[Union[Paragraph, Table]]:
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def escape_html(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def para_alignment_name(p: Paragraph) -> str:
    if p.alignment == 1:
        return "center"
    if p.alignment == 2:
        return "right"
    return "left"


# ---------- Renderers: HTML ----------


def paragraph_to_html(p: Paragraph) -> str:
    if not p.runs:
        return "<p></p>"
    parts = []
    for r in p.runs:
        t = escape_html(r.text)
        if not t:
            continue
        if r.bold:
            t = f"<b>{t}</b>"
        if r.italic:
            t = f"<i>{t}</i>"
        parts.append(t)
    align = para_alignment_name(p)
    style = f' style="text-align:{align}"' if align != "left" else ""
    return f"<p{style}>" + "".join(parts) + "</p>"


def table_to_html(tbl: Table) -> str:
    rows_html = []
    for row in tbl.rows:
        cells_html = []
        for cell in row.cells:
            cell_html = "".join(paragraph_to_html(p) for p in cell.paragraphs)
            cells_html.append(f"<td>{cell_html}</td>")
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    return (
        "<table border='1' cellpadding='4' cellspacing='0' style='border-collapse:collapse;width:100%'>"
        + "".join(rows_html)
        + "</table>"
    )


def blocks_to_html(blocks: List[Union[Paragraph, Table]]) -> str:
    html_parts = []
    for b in blocks:
        if isinstance(b, Paragraph):
            html_parts.append(paragraph_to_html(b))
        elif isinstance(b, Table):
            html_parts.append(table_to_html(b))
    return "\n".join(html_parts)


# ---------- Renderers: JSON ----------


def paragraph_to_json(p: Paragraph) -> Dict[str, Any]:
    runs = []
    for r in p.runs:
        if r.text:
            runs.append(
                {
                    "text": r.text,
                    "bold": bool(r.bold),
                    "italic": bool(r.italic),
                }
            )
    return {"type": "paragraph", "alignment": para_alignment_name(p), "runs": runs}


def table_to_json(tbl: Table) -> Dict[str, Any]:
    grid = []
    for row in tbl.rows:
        row_cells = []
        for cell in row.cells:
            row_cells.append(
                {"paragraphs": [paragraph_to_json(p) for p in cell.paragraphs]}
            )
        grid.append(row_cells)
    return {"type": "table", "rows": grid}


def blocks_to_json(blocks: List[Union[Paragraph, Table]]) -> List[Dict[str, Any]]:
    out = []
    for b in blocks:
        if isinstance(b, Paragraph):
            out.append(paragraph_to_json(b))
        elif isinstance(b, Table):
            out.append(table_to_json(b))
    return out


# ---------- Finders for sections and tables ----------


def normalize(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()


def is_title(p: Paragraph, pattern: re.Pattern) -> bool:
    return bool(pattern.match(normalize(p.text)))


def slice_blocks_between(
    doc_blocks: List[Union[Paragraph, Table]], start_idx: int, end_idx: Optional[int]
) -> List[Union[Paragraph, Table]]:
    return (
        doc_blocks[start_idx:end_idx] if end_idx is not None else doc_blocks[start_idx:]
    )


def find_section_index(
    doc_blocks: List[Union[Paragraph, Table]], title_regex: str
) -> Optional[int]:
    pat = re.compile(title_regex, flags=re.IGNORECASE)
    for i, b in enumerate(doc_blocks):
        if isinstance(b, Paragraph) and is_title(b, pat):
            return i
    return None

def find_next_section_start(
    doc_blocks: List[Union[Paragraph, Table]], from_idx: int, stop_regex: str
) -> Optional[int]:
    pat = re.compile(stop_regex, flags=re.IGNORECASE)
    for i in range(from_idx + 1, len(doc_blocks)):
        b = doc_blocks[i]
        if isinstance(b, Paragraph) and is_title(b, pat):
            return i
    return None


def extract_section_by_title(
    doc_blocks: List[Union[Paragraph, Table]],
    title_regex: str,
    next_titles_regex: List[str],
) -> List[Union[Paragraph, Table]]:
    start = find_section_index(doc_blocks, title_regex)
    if start is None:
        return []
    end_candidates = []
    for nxt in next_titles_regex:
        idx = find_next_section_start(doc_blocks, start, nxt)
        if idx is not None:
            end_candidates.append(idx)
    end = min(end_candidates) if end_candidates else None
    return slice_blocks_between(doc_blocks, start, end)


def find_table_by_caption(
    doc_blocks: List[Union[Paragraph, Table]], caption_patterns: List[str]
) -> Optional[Table]:
    """
    Find a table by searching for a caption paragraph followed by a table.
    More flexible pattern matching with broader search context.
    """
    pats = [re.compile(p, flags=re.IGNORECASE) for p in caption_patterns]
    matches = []
    for i, b in enumerate(doc_blocks):
        if isinstance(b, Paragraph):
            text = normalize(b.text)
            # Check if any pattern matches (using search for flexibility)
            if any(pat.search(text) for pat in pats):
                # Look for table within next 5 blocks (handles spacing/page breaks)
                for j in range(i + 1, min(i + 6, len(doc_blocks))):
                    if isinstance(doc_blocks[j], Table):
                        matches.append(doc_blocks[j])
                        break
    # Return the last match (likely the actual table, not from table of contents)
    return matches[-1] if matches else None


def find_recommendation_summary_table_by_structure(
    doc_blocks: List[Union[Paragraph, Table]]
) -> Optional[Table]:
    """
    Fallback method: Find recommendation summary table by analyzing table structure.
    Looks for a table with:
    - First column containing AR numbers (1, 2, 3...)
    - Headers with keywords like "Savings", "Cost", "Payback"
    - Multiple data rows
    - Located preferably in Chapter 1 section
    """
    # First, try to find Chapter 1 section to narrow search
    chapter1_start = find_section_index(doc_blocks, r"^\s*1(\.|$|\s)|^\s*EXECUTIVE\s+SUMMARY")
    chapter2_start = find_section_index(doc_blocks, r"^\s*2(\.|$|\s)|^\s*COMPANY\s+BACKGROUND")
    
    # Search within Chapter 1 if found, otherwise search entire document
    if chapter1_start is not None and chapter2_start is not None:
        search_blocks = doc_blocks[chapter1_start:chapter2_start]
    elif chapter1_start is not None:
        search_blocks = doc_blocks[chapter1_start:]
    else:
        search_blocks = doc_blocks
    
    candidate_tables = []
    
    for block in search_blocks:
        if isinstance(block, Table):
            # Check if this table looks like a recommendation summary table
            try:
                rows = block.rows
                if len(rows) < 3:  # Need header + at least 2 data rows
                    continue
                
                # Get header row text
                header_row = rows[0]
                header_text = " ".join(cell.text.lower() for cell in header_row.cells)
                
                # Check for key header terms
                has_ar_column = "ar" in header_text or "no." in header_text
                has_savings = "savings" in header_text
                has_cost = "cost" in header_text
                has_payback = "payback" in header_text
                
                if not (has_ar_column and (has_savings or has_cost)):
                    continue
                
                # Check first column for AR numbers
                ar_numbers_found = 0
                for i, row in enumerate(rows[1:], start=1):  # Skip header
                    if i > 10:  # Don't check too many rows
                        break
                    first_cell_text = row.cells[0].text.strip()
                    # Look for numeric AR numbers (1, 2, 3, etc.)
                    if first_cell_text.isdigit() or re.match(r'^\d+$', first_cell_text):
                        ar_numbers_found += 1
                
                # If we found sequential AR numbers, this is likely the summary table
                if ar_numbers_found >= 2:
                    score = ar_numbers_found
                    # Boost score if more keywords present
                    if has_payback:
                        score += 2
                    if "implementation" in header_text or "impl" in header_text:
                        score += 1
                    if "co2" in header_text or "carbon" in header_text:
                        score += 1
                    
                    candidate_tables.append((score, block))
            except Exception:
                continue
    
    # Return table with highest score
    if candidate_tables:
        candidate_tables.sort(key=lambda x: x[0], reverse=True)
        return candidate_tables[0][1]
    
    return None


def extract_ars(
    doc_blocks: List[Union[Paragraph, Table]],
) -> List[List[Union[Paragraph, Table]]]:
    """
    Extract AR sections from Chapter 4. Each AR corresponds to an "AR No. X" section.
    Only extracts ARs that are within Chapter 4/Assessment Recommendations section.
    """
    # First, find Chapter 4 start (Assessment Recommendations section)
    # This pattern catches various forms like:
    # - "4." or "4 "
    # - "4. RECOMMENDATIONS"
    # - "ASSESSMENT RECOMMENDATIONS"
    # - "RECOMMENDATIONS" (standalone)
    chapter4_start = find_section_index(
        doc_blocks,
        r"^\s*4(\.|$|\s)|^\s*(ASSESSMENT\s+)?RECOMMENDATIONS?\b",
    )

    if chapter4_start is None:
        ic("Chapter 4 not found")
        return []
        # Fallback to old method if Chapter 4 not found - search entire document
        # ar_title_pat = re.compile(r"^\s*AR\s+No\.\s*\d+\b", flags=re.IGNORECASE)
        # ar_starts: List[int] = []
        # for i, b in enumerate(doc_blocks):
        #     if isinstance(b, Paragraph) and ar_title_pat.match(normalize(b.text)):
        #         ar_starts.append(i)

        # results: List[List[Union[Paragraph, Table]]] = []
        # for k, start in enumerate(ar_starts):
        #     next_ar = ar_starts[k + 1] if k + 1 < len(ar_starts) else None
        #     next_major = find_next_section_start(
        #         doc_blocks,
        #         start,
        #         r"^\s*(5(\.|$)|INDUSTRIAL\s+CONTROL|CONCLUSIONS?|REFERENCES?|APPENDIX)",
        #     )
        #     end_candidates = [idx for idx in [next_ar, next_major] if idx is not None]
        #     end = min(end_candidates) if end_candidates else None
        #     results.append(slice_blocks_between(doc_blocks, start, end))
        # return results

    # Find Chapter 5 or other ending sections to limit our search
    # Pattern includes ".*INDUSTRIAL\s+CONTROL.*CYBERSECURITY" to match titles like
    # "THE INDUSTRIAL CONTROL SYSTEMS CYBERSECURITY ASSESSMENT" that don't start with "INDUSTRIAL"
    chapter4_end = find_next_section_start(
        doc_blocks,
        chapter4_start,
        r"^\s*(5(\.|$|\s)|Chapter\s+5|.*INDUSTRIAL\s+CONTROL.*CYBERSECURITY|CONCLUSIONS?|REFERENCES?|APPENDIX)",
    )

    # Get Chapter 4 blocks
    chapter4_blocks = slice_blocks_between(doc_blocks, chapter4_start, chapter4_end)

    # Find all AR sections in Chapter 4 using "AR No. X" pattern
    ar_pattern = re.compile(r"^\s*AR\s+No\.\s*(\d+)\b", flags=re.IGNORECASE)
    ar_starts: List[tuple[int, int]] = []  # (global_index, ar_number)

    for i, b in enumerate(chapter4_blocks):
        if isinstance(b, Paragraph):
            text = normalize(b.text)
            match = ar_pattern.match(text)
            if match:
                ar_number = int(match.group(1))
                global_index = chapter4_start + i
                ar_starts.append((global_index, ar_number))

    # Sort by AR number to ensure proper order
    ar_starts.sort(key=lambda x: x[1])

    # Extract content for each AR section
    results: List[List[Union[Paragraph, Table]]] = []
    for k, (start_idx, _) in enumerate(ar_starts):
        # Find the end of this AR section
        if k + 1 < len(ar_starts):
            # Next AR starts
            end_idx = ar_starts[k + 1][0]
        else:
            # Last AR - goes to end of Chapter 4
            end_idx = (
                chapter4_start + len(chapter4_blocks)
                if chapter4_end is None
                else chapter4_end
            )

        # Extract the blocks for this AR
        ar_blocks = slice_blocks_between(doc_blocks, start_idx, end_idx)
        results.append(ar_blocks)

    return results


def extract_ar_summaries(
    doc_blocks: List[Union[Paragraph, Table]],
) -> List[Union[Paragraph, Table]]:
    """
    Extract only the AR summary bullet points from Chapter 1.4 Summary of Best Practices and Assessment Recommendations.
    This function extracts the specific bullet point descriptions for each AR, not the full chapter content.

    Args:
        doc_blocks: List of document blocks (paragraphs and tables)

    Returns:
        List of blocks containing only the AR summary bullet points
    """
    # Extract Chapter 1.4 section
    sec_14_blocks = extract_section_by_title(
        doc_blocks,
        r"^\s*1\.4\s+Summary\s+of\s+Best\s+Practices\s+and\s+Assessment\s+Recommendations\b",
        [
            r"^\s*2(\.|$)",  # Next chapter
            r"^\s*COMPANY\s+BACKGROUND",
            r"^\s*Chapter\s+2",
            r"^\s*Section\s+2",
        ],
    )

    if not sec_14_blocks:
        # Try alternative patterns for Chapter 1.4
        sec_14_blocks = extract_section_by_title(
            doc_blocks,
            r"^\s*Summary\s+of\s+Best\s+Practices\s+and\s+Assessment\s+Recommendations\b",
            [
                r"^\s*2(\.|$)",  # Next chapter
                r"^\s*COMPANY\s+BACKGROUND",
                r"^\s*Chapter\s+2",
                r"^\s*Section\s+2",
            ],
        )

    if not sec_14_blocks:
        return []

    # Extract only AR summary bullet points
    ar_summary_blocks = []
    found_ar_summaries_start = False
    current_ar_number = None

    for i, block in enumerate(sec_14_blocks):
        if isinstance(block, Paragraph):
            text = normalize(block.text)

            # Skip Table 1.3 captions and stop processing after them
            table_13_patterns = [
                r"^\s*Table\s*1[.-]3\b.*Recommendation Summary Table",
                r"^\s*Table\s*1[.-]3\b.*Assessment Recommendation Summary Table",
            ]
            if any(
                re.match(pattern, text, re.IGNORECASE) for pattern in table_13_patterns
            ):
                break  # Stop processing once we hit Table 1.3

            # Look for the trigger phrase that indicates AR summaries are starting
            if re.search(
                r"ARs?\s+are\s+summarized\s+as\s+follows", text, re.IGNORECASE
            ):
                found_ar_summaries_start = True
                continue

            # Only start collecting after we find the "ARs are summarized as follows" phrase
            if not found_ar_summaries_start:
                continue

            # Look for AR number patterns that indicate start of a new AR summary
            ar_pattern = re.compile(r"^\s*AR\s+No\.?\s*(\d+)\s*[–\-]", re.IGNORECASE)
            ar_match = ar_pattern.match(text)
            if ar_match:
                current_ar_number = ar_match.group(1)
                ar_summary_blocks.append(block)  # Include the AR title
                continue

            # If we have an AR number and we're collecting summaries
            if current_ar_number and found_ar_summaries_start:
                # Look for bullet points or descriptions that follow the AR title
                if text.strip():
                    # Check if this looks like an AR description
                    # Bullet points typically start with • or similar, or are descriptive text
                    if (
                        text.startswith("•")
                        or text.startswith("-")
                        or text.startswith("*")
                        or text.startswith("o ")  # Sometimes bullet points are "o "
                        or (
                            len(text) > 10
                            and not ar_pattern.match(text)
                            and not text.isupper()
                        )
                    ):
                        ar_summary_blocks.append(block)

            # Stop if we hit a major section change or start of detailed content
            if (
                (text.isupper() and len(text) > 15)
                or re.match(r"^\s*2(\.|$)", text)
                or "COMPANY BACKGROUND" in text.upper()
                or "GENERAL FACILITY BACKGROUND" in text.upper()
            ):
                break

        elif isinstance(block, Table):
            # Skip all tables in this section
            continue

    return ar_summary_blocks


# ---------- Main extraction with output switch ----------


def build_outputs(blocks: List[Union[Paragraph, Table]], output: str) -> Dict[str, Any]:
    # Sections - Updated patterns to match actual document structure
    sec_11 = extract_section_by_title(
        blocks,
        r"^\s*General\s+Information\b",
        [
            r"^\s*Annual\s+Energy\s+Usages\s+and\s+Costs\b",
            r"^\s*Carbon\s+Footprint\b",
            r"^\s*Summary\s+of\s+Best\s+Practices",
        ],
    )
    sec_12 = extract_section_by_title(
        blocks,
        r"^\s*Annual\s+Energy\s+Usages\s+and\s+Costs\b",
        [r"^\s*Carbon\s+Footprint\b", r"^\s*Summary\s+of\s+Best\s+Practices"],
    )
    sec_13 = extract_section_by_title(
        blocks,
        r"^\s*Carbon\s+Footprint\b",
        [r"^\s*Summary\s+of\s+Best\s+Practices", r"^\s*COMPANY\s+BACKGROUND"],
    )

    # Table 1.3/1-3 caption (with flexible patterns and structural fallback)
    # Try caption-based search first with more flexible patterns
    # Patterns ordered from most specific to most general
    rec_tbl = find_table_by_caption(
        blocks,
        [
            r"Table\s*1[.-]3.*Assessment.*Recommendation.*Summary",  # Most specific
            r"Table\s*1[.-]3.*Recommendation.*Summary",  # Table 1.3 Recommendation Summary
            r"Table\s*1[.-]3.*Summary\s+Table",  # Table 1.3 Summary Table
            r"^Assessment\s+Recommendation\s+Summary\s+Table",  # Without table number
            r"^Recommendation\s+Summary\s+Table",  # Standalone recommendation summary
        ],
    )
    
    # If caption-based search failed, try structural identification
    if rec_tbl is None:
        rec_tbl = find_recommendation_summary_table_by_structure(blocks)

    # ARs
    ar_blocks_list = extract_ars(blocks)

    # AR Summaries from Chapter 1.4
    ar_summary_blocks = extract_ar_summaries(blocks)

    # Extract links from ARs
    ar_links = {}
    try:
        import logging
        logging.info(f"Attempting to extract links from {len(ar_blocks_list)} AR(s)")
        ar_links = extract_links_from_all_ars("", ar_blocks_list)  # docx_path not needed as we pass blocks directly
        link_stats = get_link_statistics(ar_links)
        logging.info(f"Link extraction complete: {link_stats}")
        ic(f"Extracted links from ARs: {link_stats}")
    except Exception as e:
        import logging
        import traceback
        logging.error(f"Error extracting links from ARs: {e}")
        logging.error(traceback.format_exc())
        ic(f"Error extracting links from ARs: {e}")
        ar_links = {}

    if output == "json":
        return {
            "general_information": blocks_to_json(sec_11),
            "annual_energy_usages_and_costs": blocks_to_json(sec_12),
            "carbon_footprint": blocks_to_json(sec_13),
            "recommendation_summary_table": (
                table_to_json(rec_tbl) if rec_tbl else None
            ),
            "ar_summary": blocks_to_json(ar_summary_blocks),
            "assessment_recommendations": [blocks_to_json(b) for b in ar_blocks_list],
            "ar_links": ar_links,
        }
    # default: HTML
    return {
        "general_information": blocks_to_html(sec_11),
        "annual_energy_usages_and_costs": blocks_to_html(sec_12),
        "carbon_footprint": blocks_to_html(sec_13),
        "recommendation_summary_table": (table_to_html(rec_tbl) if rec_tbl else ""),
        "ar_summary": blocks_to_html(ar_summary_blocks),
        "assessment_recommendations": [blocks_to_html(b) for b in ar_blocks_list],
        "ar_links": ar_links,
    }


def write_artifacts(payload: Dict[str, Any], output: str) -> None:
    """
    Save files to disk for inspection.
    """
    import os
    import shutil

    if output == "json":
        with open("extracted_sections.json", "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
    else:
        # Clear and recreate EXTRACTED_HTML folder
        html_folder = "EXTRACTED_HTML"
        if os.path.exists(html_folder):
            shutil.rmtree(html_folder)
        os.makedirs(html_folder, exist_ok=True)

        def write(name: str, content: str):
            file_path = os.path.join(html_folder, name)
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)

        write("general_information.html", payload["general_information"])
        write(
            "annual_energy_usages_and_costs.html",
            payload["annual_energy_usages_and_costs"],
        )
        write("carbon_footprint.html", payload["carbon_footprint"])
        write(
            "recommendation_summary_table.html", payload["recommendation_summary_table"]
        )
        write("ar_summary.html", payload["ar_summary"])
        for i, html in enumerate(payload["assessment_recommendations"], start=1):
            write(f"AR_{i:02d}.html", html)


def extract_itac_report(
    docx_path: str = DOCX_PATH, output: str = "html", save_files: bool = True
) -> Dict[str, Any]:
    """
    output: "html" or "json"
    save_files: write artifacts to disk if True
    """
    assert output in {"html", "json"}, "output must be 'html' or 'json'"
    doc = Document(docx_path)
    blocks = list(iter_block_items(doc))
    data = build_outputs(blocks, output=output)
    if save_files:
        write_artifacts(data, output=output)
    return data


def extract_general_info_fields(general_info_html: str) -> Dict[str, Union[str, float]]:
    """
    Extract specific fields from the general information HTML table.

    Args:
        general_info_html: HTML string containing the general information table

    Returns:
        Dictionary with extracted field names as keys and their values as values.
        All values are converted to float except principal_product which remains as string.
    """
    import re

    from bs4 import BeautifulSoup

    def extract_numeric_value(value_str: str) -> float:
        """Extract numeric value from a string, handling millions/billions and removing units, currency symbols, and commas."""
        # Convert to lowercase for easier matching
        value_lower = value_str.lower()

        # Remove common currency symbols and commas
        clean_str = re.sub(r"[$,]", "", value_str)

        # Find all numbers (including decimals) in the string
        numbers = re.findall(r"\d+\.?\d*", clean_str)
        if not numbers:
            return 0.0

        base_number = float(numbers[0])

        # Check for scale multipliers
        # Use more specific patterns to avoid false matches (e.g., 'k' in 'tanks', 'm' in 'employees')
        if "billion" in value_lower or re.search(r'\d+\.?\d*\s*b\b', value_lower):
            return base_number * 1_000_000_000
        elif "million" in value_lower or re.search(r'\d+\.?\d*\s*m\b', value_lower):
            return base_number * 1_000_000
        elif "thousand" in value_lower or re.search(r'\d+\.?\d*\s*k\b', value_lower):
            return base_number * 1_000
        else:
            return base_number

    # Parse the HTML
    soup = BeautifulSoup(general_info_html, "html.parser")

    # Initialize the result dictionary
    extracted_fields = {}

    # Find the table containing the general information
    table = soup.find("table")
    if not table:
        return extracted_fields

    # Extract data from table rows
    for row in table.find_all("tr"):
        cells = row.find_all("td")
        for cell in cells:
            cell_text = cell.get_text(strip=True)
            if ":" in cell_text:
                # Split on the first colon to separate field name and value
                parts = cell_text.split(":", 1)
                if len(parts) == 2:
                    field_name = parts[0].strip()
                    field_value = parts[1].strip()

                    # Normalize field names to consistent keys
                    field_mapping = {
                        "SIC. No.": "sic_no",
                        "SIC No.": "sic_no",
                        "SIC No": "sic_no",
                        "NAICS Code": "naics_code",
                        "Principal Product": "principal_product",
                        "Principal Products": "principal_products",
                        "No. of Employees": "no_of_employees",
                        "Number of Employees": "no_of_employees",
                        "Total Facility Area": "total_facility_area",
                        "Operating Hours": "operating_hours",
                        "Annual Production": "annual_production",
                        "Annual Sales": "annual_sales",
                        "Value per Finished Product": "value_per_finished_product",
                        "Total Energy Usage": "total_energy_usage",
                        "Total Utility Cost": "total_utility_cost",
                        "No. of Assessment Recommendations": "no_of_assessment_recommendations",
                    }

                    # Map to standardized key or use original field name
                    standardized_key = field_mapping.get(
                        field_name,
                        field_name.lower().replace(" ", "_").replace(".", ""),
                    )

                    # Convert to appropriate type
                    if standardized_key in ["principal_product", "principal_products"]:
                        extracted_fields[standardized_key] = field_value
                    else:
                        extracted_fields[standardized_key] = extract_numeric_value(
                            field_value
                        )

    return extracted_fields


def extract_energy_usage(annual_energy_html: str) -> Dict[str, Any]:
    """
    Extract energy usage data from the annual energy usages and costs HTML.

    Args:
        annual_energy_html: HTML string containing the annual energy usages and costs table

    Returns:
        Dictionary with period information and energy usage data
    """
    import re

    from bs4 import BeautifulSoup

    def extract_period_from_text(text: str) -> Dict[str, str]:
        """Extract start and end period from descriptive text."""
        # Look for patterns like "between September 2023 and August 2024"
        period_pattern = r"between\s+(\w+\s+\d{4})\s+and\s+(\w+\s+\d{4})"
        match = re.search(period_pattern, text, re.IGNORECASE)
        if match:
            return {"start": match.group(1), "end": match.group(2)}

        # Alternative pattern: "from X to Y" or "X - Y"
        alt_pattern = r"(?:from\s+)?(\w+\s+\d{4})(?:\s+(?:to|-)\s+)(\w+\s+\d{4})"
        match = re.search(alt_pattern, text, re.IGNORECASE)
        if match:
            return {"start": match.group(1), "end": match.group(2)}

        return {"start": "", "end": ""}

    def parse_usage_cell(usage_text: str) -> Dict[str, float]:
        """Parse usage cell that may contain multiple values with different units."""
        usage_dict = {}

        # Find all patterns like "649,680 kWh/yr" or "(2,217 MMBTU/yr)"
        patterns = re.findall(
            r"[\(]?([0-9,]+\.?[0-9]*)\s+([A-Za-z/]+)[\)]?", usage_text
        )

        for value_str, unit in patterns:
            # Clean up the value string and convert to float
            clean_value = re.sub(r"[,\s]", "", value_str)
            try:
                value = float(clean_value)
                usage_dict[unit] = value
            except ValueError:
                continue

        return usage_dict

    def parse_cost_cell(cost_text: str) -> float:
        """Parse cost cell and return numeric value."""
        # Remove currency symbols, commas, and /yr
        clean_cost = re.sub(r"[\$,/yr\s]", "", cost_text)
        # Extract numeric value
        numbers = re.findall(r"\d+\.?\d*", clean_cost)
        if numbers:
            return float(numbers[0])
        return 0.0

    def parse_unit_cost_cell(
        unit_cost_text: str,
    ) -> Optional[Dict[str, Union[float, str]]]:
        """Parse unit cost cell like '$0.102/kWh' or '$4.522/kW'."""
        if unit_cost_text.strip() in ["-", ""]:
            return None

        # Pattern for $X.XX/unit
        pattern = r"\$([0-9,]+\.?[0-9]*)/([A-Za-z]+)"
        match = re.search(pattern, unit_cost_text)
        if match:
            amount = float(re.sub(r"[,]", "", match.group(1)))
            unit = match.group(2)
            return {"amount": amount, "unit": unit}
        return None

    # Parse the HTML
    soup = BeautifulSoup(annual_energy_html, "html.parser")

    # Initialize result structure
    result = {"period": {"start": "", "end": ""}, "data": []}

    # Extract period information from paragraph text
    paragraphs = soup.find_all("p")
    for p in paragraphs:
        text = p.get_text()
        period_info = extract_period_from_text(text)
        if period_info["start"] and period_info["end"]:
            result["period"] = period_info
            break

    # Find the energy usage table
    table = soup.find("table")
    if not table:
        return result

    # Process table rows (skip header)
    rows = table.find_all("tr")
    if len(rows) <= 1:  # No data rows
        return result

    for row in rows[1:]:  # Skip header row
        cells = row.find_all("td")
        if len(cells) < 4:  # Should have Type, Usage, Cost, Unit Cost
            continue

        # Extract data from each cell
        energy_type_raw = cells[0].get_text(strip=True).replace("**", "").strip()
        usage_text = cells[1].get_text()
        cost_text = cells[2].get_text()
        unit_cost_text = cells[3].get_text()

        # Map energy types to programming-oriented field names
        type_mapping = {
            "Electrical Energy": "electrical_energy",
            "Electrical Demand": "electrical_demand",
            "Electric Energy": "electrical_energy",
            "Electric Demand": "electrical_demand",
            "Electricity": "electrical_energy",
            "Demand Charge": "electrical_demand",  # TODO: Verify this is correct
            "Demand": "electrical_demand",
            "Natural Gas": "natural_gas",
            "Propane": "propane_gas",
            "Propane Gas": "propane_gas",
            "Steam": "steam",
            "Water": "water",
            "Compressed Air": "compressed_air",
            "Total Utility": "total_utility",
            "TotalUtility": "total_utility",
            "Total": "total_utility",
            "Fuel Oil": "fuel_oil",
            "Heating Oil": "heating_oil",
            "Diesel": "diesel",
            "Gasoline": "gasoline",
            "Coal": "coal",
            "Biomass": "biomass",
            "Solar": "solar",
            "Wind": "wind",
            "Geothermal": "geothermal",
            "Chilled Water": "chilled_water",
            "Hot Water": "hot_water",
        }

        # Get standardized type name or create one from raw name
        energy_type = type_mapping.get(
            energy_type_raw,
            energy_type_raw.lower()
            .replace(" ", "_")
            .replace("-", "_")
            .replace("&", "and")
            .replace("/", "_"),
        )

        # Parse the data
        usage_data = parse_usage_cell(usage_text)
        cost_value = parse_cost_cell(cost_text)
        unit_cost_data = parse_unit_cost_cell(unit_cost_text)

        # Create entry
        entry = {
            "type": energy_type,
            "usage": usage_data,
            "cost": cost_value,
            "unit_cost": unit_cost_data,
        }

        result["data"].append(entry)

    return result


if __name__ == "__main__":
    # HTML run
    html_out = extract_itac_report(DOCX_PATH, output="html", save_files=True)
    print(
        "HTML extraction complete.",
        len(html_out["assessment_recommendations"]),
        "AR sections",
    )

    # JSON run
    json_out = extract_itac_report(DOCX_PATH, output="json", save_files=True)
    print(
        "JSON extraction complete.",
        len(json_out["assessment_recommendations"]),
        "AR sections",
    )

    # Extract general information fields
    general_info_fields = extract_general_info_fields(html_out["general_information"])
    print("\nExtracted General Information Fields:")
    for key, value in general_info_fields.items():
        print(f"  {key}: {value}")

    # Extract energy usage data
    energy_usage_data = extract_energy_usage(html_out["annual_energy_usages_and_costs"])
    print(f"\nExtracted Energy Usage Data:")
    print(
        f"  Period: {energy_usage_data['period']['start']} to {energy_usage_data['period']['end']}"
    )
    print(f"  Number of energy types: {len(energy_usage_data['data'])}")
    for item in energy_usage_data["data"]:
        print(f"    - {item['type']}: {item['usage']} (Cost: ${item['cost']:.2f})")
